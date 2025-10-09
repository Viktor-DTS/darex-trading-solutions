#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Скрипт для аналізу Word документа та витягнення всіх стилів та структури
"""

import os
import sys
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json

def analyze_word_document(doc_path):
    """Аналізує Word документ та витягає всі стилі та структуру"""
    
    if not os.path.exists(doc_path):
        print(f"Файл не знайдено: {doc_path}")
        return None
    
    try:
        doc = Document(doc_path)
        analysis = {
            "document_info": {
                "filename": os.path.basename(doc_path),
                "total_paragraphs": len(doc.paragraphs),
                "total_tables": len(doc.tables)
            },
            "paragraphs": [],
            "tables": [],
            "styles_used": set()
        }
        
        print("=== АНАЛІЗ WORD ДОКУМЕНТА ===")
        print(f"Файл: {doc_path}")
        print(f"Кількість абзаців: {len(doc.paragraphs)}")
        print(f"Кількість таблиць: {len(doc.tables)}")
        print("\n" + "="*50)
        
        # Аналіз абзаців
        for i, paragraph in enumerate(doc.paragraphs):
            if paragraph.text.strip():  # Тільки непусті абзаци
                para_info = {
                    "index": i,
                    "text": paragraph.text.strip(),
                    "alignment": str(paragraph.alignment) if paragraph.alignment else "None",
                    "runs": []
                }
                
                # Аналіз кожного run (частини тексту з різними стилями)
                for run in paragraph.runs:
                    if run.text.strip():
                        run_info = {
                            "text": run.text.strip(),
                            "bold": run.bold,
                            "italic": run.italic,
                            "underline": run.underline,
                            "font_name": run.font.name,
                            "font_size": str(run.font.size) if run.font.size else "None",
                            "font_color": str(run.font.color.rgb) if run.font.color and run.font.color.rgb else "None"
                        }
                        para_info["runs"].append(run_info)
                        
                        # Збираємо унікальні стилі
                        if run.font.name:
                            analysis["styles_used"].add(f"Font: {run.font.name}")
                        if run.font.size:
                            analysis["styles_used"].add(f"Size: {run.font.size}")
                
                analysis["paragraphs"].append(para_info)
                
                # Виводимо інформацію про абзац
                print(f"\nАБЗАЦ {i+1}:")
                print(f"Текст: {paragraph.text.strip()}")
                print(f"Вирівнювання: {paragraph.alignment}")
                
                if para_info["runs"]:
                    print("Стилі тексту:")
                    for j, run in enumerate(para_info["runs"]):
                        print(f"  Run {j+1}: '{run['text']}'")
                        print(f"    Шрифт: {run['font_name']}")
                        print(f"    Розмір: {run['font_size']}")
                        print(f"    Жирний: {run['bold']}")
                        print(f"    Курсив: {run['italic']}")
                        print(f"    Підкреслений: {run['underline']}")
                        print(f"    Колір: {run['font_color']}")
        
        # Аналіз таблиць
        for i, table in enumerate(doc.tables):
            table_info = {
                "index": i,
                "rows": len(table.rows),
                "columns": len(table.columns),
                "cells": []
            }
            
            print(f"\nТАБЛИЦЯ {i+1}:")
            print(f"Рядків: {len(table.rows)}")
            print(f"Колонок: {len(table.columns)}")
            
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    if cell.text.strip():
                        cell_info = {
                            "row": row_idx,
                            "column": col_idx,
                            "text": cell.text.strip(),
                            "paragraphs": []
                        }
                        
                        # Аналіз абзаців в комірці
                        for para in cell.paragraphs:
                            if para.text.strip():
                                para_info = {
                                    "text": para.text.strip(),
                                    "alignment": str(para.alignment) if para.alignment else "None",
                                    "runs": []
                                }
                                
                                for run in para.runs:
                                    if run.text.strip():
                                        run_info = {
                                            "text": run.text.strip(),
                                            "bold": run.bold,
                                            "italic": run.italic,
                                            "font_name": run.font.name,
                                            "font_size": str(run.font.size) if run.font.size else "None"
                                        }
                                        para_info["runs"].append(run_info)
                                
                                cell_info["paragraphs"].append(para_info)
                        
                        table_info["cells"].append(cell_info)
                        print(f"  [{row_idx+1},{col_idx+1}]: {cell.text.strip()}")
            
            analysis["tables"].append(table_info)
        
        # Виводимо унікальні стилі
        print(f"\nУНІКАЛЬНІ СТИЛІ:")
        for style in sorted(analysis["styles_used"]):
            print(f"  {style}")
        
        # Зберігаємо детальний аналіз в JSON файл
        output_file = "word_analysis.json"
        with open(output_file, 'w', encoding='utf-8') as f:
            json.dump(analysis, f, ensure_ascii=False, indent=2, default=str)
        
        print(f"\nДетальний аналіз збережено в файл: {output_file}")
        
        return analysis
        
    except Exception as e:
        print(f"Помилка при аналізі документа: {e}")
        return None

def main():
    # Шлях до Word документа
    doc_path = r"C:\dts-service\Наряд на работу V2.51.0У ТО.docx"
    
    print("Запуск аналізу Word документа...")
    analysis = analyze_word_document(doc_path)
    
    if analysis:
        print("\n✅ Аналіз завершено успішно!")
        print("Тепер ви можете скопіювати вивід вище та передати мені для створення точного шаблону.")
    else:
        print("❌ Помилка при аналізі документа")

if __name__ == "__main__":
    main()
